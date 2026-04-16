import csv
import io
import json
import xml.etree.ElementTree as ET
from pathlib import Path

import openpyxl


def _make_result(cell_values, sheet_names=None, headers=None, sample_rows=None, truncated=False):
    return {
        "body_text": " ".join(cell_values),
        "sheet_names": sheet_names or [],
        "headers": headers or {},
        "sample_rows": sample_rows or {},
        "truncated": truncated,
    }


def read_xlsx(path: str | Path, max_rows: int = 10000) -> dict:
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    cell_values = []
    truncated = False
    headers_map = {}
    sample_rows_map = {}

    for name in wb.sheetnames:
        ws = wb[name]
        rows_iter = ws.iter_rows(values_only=True)

        first_row = next(rows_iter, None)
        if first_row is None:
            headers_map[name] = []
            sample_rows_map[name] = []
            continue

        headers = [str(c) if c is not None else "" for c in first_row]
        headers_map[name] = headers

        for h in headers:
            if h:
                cell_values.append(h)

        sample = []
        for i, row in enumerate(rows_iter):
            if i >= max_rows:
                truncated = True
                break
            row_dict = {}
            for j, val in enumerate(row):
                if val is not None:
                    key = headers[j] if j < len(headers) else f"col_{j}"
                    row_dict[key] = val
                    if isinstance(val, str) and val.strip():
                        cell_values.append(val.strip())
            if row_dict:
                sample.append(row_dict)
        sample_rows_map[name] = sample

    wb.close()
    return _make_result(cell_values, sheet_names=wb.sheetnames, headers=headers_map,
                        sample_rows=sample_rows_map, truncated=truncated)


def read_pdf(path: str | Path, max_pages: int = 500) -> dict:
    import pymupdf
    doc = pymupdf.open(str(path))
    values = []
    truncated = False
    for i, page in enumerate(doc):
        if i >= max_pages:
            truncated = True
            break
        text = page.get_text().strip()
        if text:
            values.extend(text.split())
    doc.close()
    if not values:
        return None
    return _make_result(values, truncated=truncated)


def read_docx(path: str | Path) -> dict:
    import docx
    doc = docx.Document(str(path))
    values = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            values.extend(text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    values.append(text)
    if not values:
        return None
    return _make_result(values)


def read_text(path: str | Path, max_chars: int = 500000) -> dict:
    raw = Path(path).read_text(errors="replace")
    truncated = len(raw) > max_chars
    text = raw[:max_chars]
    values = text.split()
    if not values:
        return None
    return _make_result(values, truncated=truncated)


def read_csv_file(path: str | Path, max_chars: int = 500000) -> dict:
    raw = Path(path).read_text(errors="replace")
    truncated = len(raw) > max_chars
    text = raw[:max_chars]
    reader = csv.reader(io.StringIO(text))
    values = []
    headers = {}
    first_row = None
    for i, row in enumerate(reader):
        if i == 0:
            first_row = row
            headers["sheet1"] = row
            values.extend(row)
        else:
            values.extend(c for c in row if c.strip())
    if not values:
        return None
    return _make_result(values, headers=headers, truncated=truncated)


def read_xml(path: str | Path) -> dict:
    try:
        tree = ET.parse(str(path))
    except ET.ParseError:
        return None
    values = []
    for elem in tree.getroot().iter():
        if elem.text and elem.text.strip():
            values.append(elem.text.strip())
        for v in elem.attrib.values():
            if v.strip():
                values.append(v.strip())
    if not values:
        return None
    return _make_result(values[:2000])


def read_json_file(path: str | Path) -> dict:
    text = Path(path).read_text(errors="replace")[:50000]
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return None

    values = []
    def extract(obj, depth=0):
        if depth > 5:
            return
        if isinstance(obj, dict):
            for k, v in obj.items():
                values.append(str(k))
                extract(v, depth + 1)
        elif isinstance(obj, list):
            for item in obj[:50]:
                extract(item, depth + 1)
        elif isinstance(obj, str) and obj.strip():
            values.append(obj.strip())

    extract(data)
    if not values:
        return None
    return _make_result(values[:2000])


CONTENT_READERS = {
    ".xlsx": read_xlsx,
    ".xls": read_xlsx,
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".doc": read_docx,
    ".txt": read_text,
    ".md": read_text,
    ".csv": read_csv_file,
    ".xml": read_xml,
    ".json": read_json_file,
    ".html": read_text,
}


def extract_text(path: str | Path) -> dict | None:
    path = Path(path)
    ext = path.suffix.lower()
    reader = CONTENT_READERS.get(ext)
    if reader is None:
        return None
    try:
        return reader(path)
    except Exception:
        return None
